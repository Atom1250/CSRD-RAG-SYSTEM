"""
Text processing service for document text extraction, chunking, and preprocessing
"""
import re
import logging
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple
from io import BytesIO

import PyPDF2
import pdfplumber
from docx import Document as DocxDocument
from sqlalchemy.orm import Session

from app.models.database import Document, TextChunk, DocumentType, ProcessingStatus
from app.models.schemas import TextChunkCreate, TextChunkResponse
from app.core.config import settings

# Conditional import for vector service to avoid dependency issues during testing
try:
    from app.services.vector_service import embedding_service
    VECTOR_SERVICE_AVAILABLE = True
except ImportError as e:
    # Logger is defined below, so we'll set a flag and log later
    embedding_service = None
    VECTOR_SERVICE_AVAILABLE = False
    _vector_import_error = str(e)

logger = logging.getLogger(__name__)


class TextProcessingConfig:
    """Configuration for text processing operations"""
    
    # Default chunk sizes
    DEFAULT_CHUNK_SIZE = 1000
    DEFAULT_CHUNK_OVERLAP = 200
    MIN_CHUNK_SIZE = 100
    MAX_CHUNK_SIZE = 5000
    
    # Text cleaning patterns
    WHITESPACE_PATTERN = re.compile(r'\s+')
    SPECIAL_CHARS_PATTERN = re.compile(r'[^\w\s\.\,\!\?\;\:\-\(\)\[\]\{\}\"\'\/\\]')
    MULTIPLE_NEWLINES_PATTERN = re.compile(r'\n{3,}')
    
    # PDF extraction settings
    PDF_EXTRACTION_FALLBACK = True
    PDF_MIN_CONFIDENCE = 0.5


class TextExtractionError(Exception):
    """Custom exception for text extraction errors"""
    pass


class TextChunkingError(Exception):
    """Custom exception for text chunking errors"""
    pass


class TextProcessingService:
    """Service for text extraction, chunking, and preprocessing"""
    
    def __init__(self, db: Session):
        self.db = db
        self.config = TextProcessingConfig()
        
        # Log vector service availability
        if not VECTOR_SERVICE_AVAILABLE:
            logger.warning(f"Vector service not available: {_vector_import_error if '_vector_import_error' in globals() else 'Unknown error'}")
    
    def extract_text_from_document(self, document: Document) -> str:
        """
        Extract text from a document based on its type
        
        Args:
            document: Document database model instance
            
        Returns:
            str: Extracted text content
            
        Raises:
            TextExtractionError: If text extraction fails
        """
        file_path = Path(document.file_path)
        
        if not file_path.exists():
            raise TextExtractionError(f"Document file not found: {file_path}")
        
        try:
            if document.document_type == DocumentType.PDF:
                return self._extract_text_from_pdf(file_path)
            elif document.document_type == DocumentType.DOCX:
                return self._extract_text_from_docx(file_path)
            elif document.document_type == DocumentType.TXT:
                return self._extract_text_from_txt(file_path)
            else:
                raise TextExtractionError(f"Unsupported document type: {document.document_type}")
                
        except Exception as e:
            logger.error(f"Text extraction failed for document {document.id}: {str(e)}")
            raise TextExtractionError(f"Failed to extract text: {str(e)}")
    
    def _extract_text_from_pdf(self, file_path: Path) -> str:
        """
        Extract text from PDF using pdfplumber with PyPDF2 fallback
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            str: Extracted text content
        """
        text_content = ""
        
        try:
            # Primary extraction using pdfplumber
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text_content += f"\n--- Page {page_num + 1} ---\n"
                            text_content += page_text + "\n"
                    except Exception as e:
                        logger.warning(f"Failed to extract text from page {page_num + 1}: {str(e)}")
                        continue
            
            # If pdfplumber extraction is insufficient, try PyPDF2 fallback
            if self.config.PDF_EXTRACTION_FALLBACK and len(text_content.strip()) < 100:
                logger.info("Using PyPDF2 fallback for PDF text extraction")
                text_content = self._extract_text_from_pdf_pypdf2(file_path)
                
        except Exception as e:
            logger.warning(f"pdfplumber extraction failed: {str(e)}, trying PyPDF2 fallback")
            text_content = self._extract_text_from_pdf_pypdf2(file_path)
        
        if not text_content.strip():
            raise TextExtractionError("No text content could be extracted from PDF")
        
        return text_content.strip()
    
    def _extract_text_from_pdf_pypdf2(self, file_path: Path) -> str:
        """
        Fallback PDF text extraction using PyPDF2
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            str: Extracted text content
        """
        text_content = ""
        
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text_content += f"\n--- Page {page_num + 1} ---\n"
                        text_content += page_text + "\n"
                except Exception as e:
                    logger.warning(f"PyPDF2 failed to extract text from page {page_num + 1}: {str(e)}")
                    continue
        
        return text_content.strip()
    
    def _extract_text_from_docx(self, file_path: Path) -> str:
        """
        Extract text from DOCX file
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            str: Extracted text content
        """
        try:
            doc = DocxDocument(file_path)
            text_content = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text.strip())
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(" | ".join(row_text))
            
            full_text = "\n".join(text_content)
            
            if not full_text.strip():
                raise TextExtractionError("No text content found in DOCX file")
            
            return full_text.strip()
            
        except Exception as e:
            raise TextExtractionError(f"Failed to extract text from DOCX: {str(e)}")
    
    def _extract_text_from_txt(self, file_path: Path) -> str:
        """
        Extract text from TXT file with encoding detection
        
        Args:
            file_path: Path to TXT file
            
        Returns:
            str: Extracted text content
        """
        encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    content = file.read()
                    if content.strip():
                        return content.strip()
            except (UnicodeDecodeError, UnicodeError):
                continue
            except Exception as e:
                raise TextExtractionError(f"Failed to read TXT file: {str(e)}")
        
        raise TextExtractionError("Could not decode TXT file with any supported encoding")
    
    def preprocess_text(self, text: str) -> str:
        """
        Clean and normalize text content
        
        Args:
            text: Raw text content
            
        Returns:
            str: Preprocessed text
        """
        if not text or not text.strip():
            return ""
        
        # Remove excessive whitespace
        text = self.config.WHITESPACE_PATTERN.sub(' ', text)
        
        # Remove excessive newlines but preserve paragraph structure
        text = self.config.MULTIPLE_NEWLINES_PATTERN.sub('\n\n', text)
        
        # Remove special characters but keep basic punctuation
        # This is conservative to preserve document structure
        text = text.replace('\x00', '')  # Remove null characters
        text = text.replace('\ufeff', '')  # Remove BOM
        
        # Normalize line endings
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        
        # Remove leading/trailing whitespace from each line
        lines = [line.strip() for line in text.split('\n')]
        text = '\n'.join(lines)
        
        return text.strip()
    
    def chunk_text(
        self, 
        text: str, 
        chunk_size: int = None, 
        chunk_overlap: int = None
    ) -> List[str]:
        """
        Split text into chunks with configurable size and overlap
        
        Args:
            text: Text to chunk
            chunk_size: Maximum characters per chunk
            chunk_overlap: Number of characters to overlap between chunks
            
        Returns:
            List[str]: List of text chunks
            
        Raises:
            TextChunkingError: If chunking parameters are invalid
        """
        if not text or not text.strip():
            return []
        
        # Use default values if not provided
        chunk_size = chunk_size or self.config.DEFAULT_CHUNK_SIZE
        chunk_overlap = chunk_overlap or self.config.DEFAULT_CHUNK_OVERLAP
        
        # Validate parameters
        if chunk_size < self.config.MIN_CHUNK_SIZE:
            raise TextChunkingError(f"Chunk size must be at least {self.config.MIN_CHUNK_SIZE}")
        if chunk_size > self.config.MAX_CHUNK_SIZE:
            raise TextChunkingError(f"Chunk size cannot exceed {self.config.MAX_CHUNK_SIZE}")
        if chunk_overlap >= chunk_size:
            raise TextChunkingError("Chunk overlap must be less than chunk size")
        if chunk_overlap < 0:
            raise TextChunkingError("Chunk overlap cannot be negative")
        
        chunks = []
        text_length = len(text)
        start = 0
        
        while start < text_length:
            # Calculate end position
            end = start + chunk_size
            
            # If this is not the last chunk, try to break at a natural boundary
            if end < text_length:
                # Look for sentence boundaries within the last 20% of the chunk
                boundary_search_start = max(start + int(chunk_size * 0.8), start + 1)
                
                # Try to find sentence end
                sentence_end = self._find_sentence_boundary(text, boundary_search_start, end)
                if sentence_end > start:
                    end = sentence_end
                else:
                    # Try to find word boundary
                    word_end = self._find_word_boundary(text, end)
                    if word_end > start:
                        end = word_end
            
            # Extract chunk
            chunk = text[start:end].strip()
            
            if chunk:  # Only add non-empty chunks
                chunks.append(chunk)
            
            # Calculate next start position with overlap
            if end >= text_length:
                break
            
            start = max(end - chunk_overlap, start + 1)
        
        return chunks
    
    def _find_sentence_boundary(self, text: str, start: int, end: int) -> int:
        """
        Find the best sentence boundary within a range
        
        Args:
            text: Full text
            start: Start position to search
            end: End position to search
            
        Returns:
            int: Position of sentence boundary, or -1 if not found
        """
        sentence_endings = ['. ', '! ', '? ', '.\n', '!\n', '?\n']
        
        best_pos = -1
        for i in range(end - 1, start - 1, -1):
            for ending in sentence_endings:
                if text[i:i + len(ending)] == ending:
                    return i + 1
        
        return best_pos
    
    def _find_word_boundary(self, text: str, position: int) -> int:
        """
        Find the nearest word boundary before the given position
        
        Args:
            text: Full text
            position: Position to search backwards from
            
        Returns:
            int: Position of word boundary
        """
        if position >= len(text):
            return len(text)
        
        # Search backwards for whitespace
        for i in range(position - 1, max(0, position - 100), -1):
            if text[i].isspace():
                return i + 1
        
        return position
    
    async def process_document_text(
        self, 
        document_id: str, 
        chunk_size: int = None, 
        chunk_overlap: int = None,
        generate_embeddings: bool = True
    ) -> List[TextChunkResponse]:
        """
        Complete text processing pipeline for a document
        
        Args:
            document_id: Document ID to process
            chunk_size: Optional chunk size override
            chunk_overlap: Optional chunk overlap override
            generate_embeddings: Whether to generate and store embeddings
            
        Returns:
            List[TextChunkResponse]: Created text chunks
            
        Raises:
            TextExtractionError: If document processing fails
        """
        # Get document from database
        document = self.db.query(Document).filter(Document.id == document_id).first()
        if not document:
            raise TextExtractionError(f"Document not found: {document_id}")
        
        # Update processing status
        document.processing_status = ProcessingStatus.PROCESSING
        self.db.commit()
        
        try:
            # Extract text
            raw_text = self.extract_text_from_document(document)
            
            # Preprocess text
            processed_text = self.preprocess_text(raw_text)
            
            # Chunk text
            chunks = self.chunk_text(processed_text, chunk_size, chunk_overlap)
            
            # Create text chunk records
            created_chunks = []
            chunks_for_embedding = []
            
            for i, chunk_content in enumerate(chunks):
                chunk_data = TextChunkCreate(
                    document_id=document_id,
                    content=chunk_content,
                    chunk_index=i
                )
                
                db_chunk = TextChunk(
                    document_id=chunk_data.document_id,
                    content=chunk_data.content,
                    chunk_index=chunk_data.chunk_index
                )
                
                self.db.add(db_chunk)
                self.db.flush()  # Get the ID
                
                chunk_response = TextChunkResponse.model_validate(db_chunk)
                created_chunks.append(chunk_response)
                
                # Prepare for embedding generation
                if generate_embeddings:
                    chunks_for_embedding.append({
                        "id": db_chunk.id,
                        "document_id": document_id,
                        "content": chunk_content,
                        "chunk_index": i,
                        "schema_elements": [],  # Will be populated by schema service
                        "created_at": db_chunk.created_at.isoformat() if db_chunk.created_at else ""
                    })
            
            # Generate and store embeddings
            if generate_embeddings and chunks_for_embedding and VECTOR_SERVICE_AVAILABLE:
                try:
                    embedding_success = await embedding_service.store_embeddings(chunks_for_embedding)
                    if embedding_success:
                        # Update database with embedding vectors
                        for i, chunk_data in enumerate(chunks_for_embedding):
                            if "embedding_vector" in chunk_data:
                                db_chunk = self.db.query(TextChunk).filter(TextChunk.id == chunk_data["id"]).first()
                                if db_chunk:
                                    db_chunk.embedding_vector = chunk_data["embedding_vector"]
                        
                        logger.info(f"Generated embeddings for {len(chunks_for_embedding)} chunks")
                    else:
                        logger.warning(f"Failed to generate embeddings for document {document_id}")
                        
                except Exception as e:
                    logger.error(f"Embedding generation failed for document {document_id}: {str(e)}")
                    # Don't fail the entire process if embedding generation fails
            elif generate_embeddings and not VECTOR_SERVICE_AVAILABLE:
                logger.warning("Embedding generation requested but vector service is not available")
            
            # Update document status to completed
            document.processing_status = ProcessingStatus.COMPLETED
            self.db.commit()
            
            logger.info(f"Successfully processed document {document_id} into {len(created_chunks)} chunks")
            return created_chunks
            
        except Exception as e:
            # Update document status to failed
            document.processing_status = ProcessingStatus.FAILED
            self.db.commit()
            
            logger.error(f"Document processing failed for {document_id}: {str(e)}")
            raise TextExtractionError(f"Document processing failed: {str(e)}")
    
    def get_document_chunks(self, document_id: str) -> List[TextChunkResponse]:
        """
        Retrieve all text chunks for a document
        
        Args:
            document_id: Document ID
            
        Returns:
            List[TextChunkResponse]: Document text chunks
        """
        chunks = (
            self.db.query(TextChunk)
            .filter(TextChunk.document_id == document_id)
            .order_by(TextChunk.chunk_index)
            .all()
        )
        
        return [TextChunkResponse.model_validate(chunk) for chunk in chunks]
    
    async def regenerate_embeddings(self, document_id: str) -> bool:
        """
        Regenerate embeddings for all chunks of a document
        
        Args:
            document_id: Document ID
            
        Returns:
            bool: True if successful, False otherwise
        """
        if not VECTOR_SERVICE_AVAILABLE:
            logger.warning("Vector service not available for embedding regeneration")
            return False
            
        try:
            chunks = self.get_document_chunks(document_id)
            if not chunks:
                logger.warning(f"No chunks found for document {document_id}")
                return False
            
            # Prepare chunks for embedding
            chunks_for_embedding = []
            for chunk in chunks:
                chunks_for_embedding.append({
                    "id": chunk.id,
                    "document_id": chunk.document_id,
                    "content": chunk.content,
                    "chunk_index": chunk.chunk_index,
                    "schema_elements": chunk.schema_elements or [],
                    "created_at": chunk.created_at.isoformat() if chunk.created_at else ""
                })
            
            # Generate and store embeddings
            embedding_success = await embedding_service.store_embeddings(chunks_for_embedding)
            
            if embedding_success:
                # Update database with new embedding vectors
                for chunk_data in chunks_for_embedding:
                    if "embedding_vector" in chunk_data:
                        db_chunk = self.db.query(TextChunk).filter(TextChunk.id == chunk_data["id"]).first()
                        if db_chunk:
                            db_chunk.embedding_vector = chunk_data["embedding_vector"]
                
                self.db.commit()
                logger.info(f"Successfully regenerated embeddings for {len(chunks)} chunks in document {document_id}")
                return True
            else:
                logger.error(f"Failed to regenerate embeddings for document {document_id}")
                return False
                
        except Exception as e:
            logger.error(f"Error regenerating embeddings for document {document_id}: {str(e)}")
            return False
    
    async def delete_document_embeddings(self, document_id: str) -> bool:
        """
        Delete all embeddings for a document
        
        Args:
            document_id: Document ID
            
        Returns:
            bool: True if successful, False otherwise
        """
        if not VECTOR_SERVICE_AVAILABLE:
            logger.warning("Vector service not available for embedding deletion")
            return False
            
        try:
            chunks = self.get_document_chunks(document_id)
            if not chunks:
                return True  # No chunks to delete
            
            chunk_ids = [chunk.id for chunk in chunks]
            
            # Delete from vector database
            delete_success = await embedding_service.delete_chunk_embeddings(chunk_ids)
            
            if delete_success:
                # Clear embedding vectors from database
                for chunk_id in chunk_ids:
                    db_chunk = self.db.query(TextChunk).filter(TextChunk.id == chunk_id).first()
                    if db_chunk:
                        db_chunk.embedding_vector = None
                
                self.db.commit()
                logger.info(f"Successfully deleted embeddings for {len(chunk_ids)} chunks in document {document_id}")
                return True
            else:
                logger.error(f"Failed to delete embeddings for document {document_id}")
                return False
                
        except Exception as e:
            logger.error(f"Error deleting embeddings for document {document_id}: {str(e)}")
            return False
    
    def get_processing_statistics(self, document_id: str) -> Dict[str, Any]:
        """
        Get processing statistics for a document
        
        Args:
            document_id: Document ID
            
        Returns:
            Dict containing processing statistics
        """
        document = self.db.query(Document).filter(Document.id == document_id).first()
        if not document:
            return {}
        
        chunks = self.get_document_chunks(document_id)
        
        if not chunks:
            return {
                "document_id": document_id,
                "processing_status": document.processing_status,
                "total_chunks": 0,
                "total_characters": 0,
                "average_chunk_size": 0,
                "embeddings_generated": 0
            }
        
        total_chars = sum(len(chunk.content) for chunk in chunks)
        avg_chunk_size = total_chars / len(chunks) if chunks else 0
        embeddings_count = sum(1 for chunk in chunks if chunk.embedding_vector)
        
        return {
            "document_id": document_id,
            "processing_status": document.processing_status,
            "total_chunks": len(chunks),
            "total_characters": total_chars,
            "average_chunk_size": round(avg_chunk_size, 2),
            "chunk_sizes": [len(chunk.content) for chunk in chunks],
            "embeddings_generated": embeddings_count,
            "embedding_coverage": round(embeddings_count / len(chunks) * 100, 2) if chunks else 0
        }